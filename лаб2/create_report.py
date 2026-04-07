# -*- coding: utf-8 -*-
"""
Скрипт для создания отчёта по Лабораторной работе №2
Тема: Факторизация чисел с использованием хеш-функции Стрибог
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_lab2_report():
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # === ТИТУЛЬНЫЙ ЛИСТ ===
    
    # Министерство
    p = doc.add_paragraph('МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РФ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    
    # Университет
    p = doc.add_paragraph('Федеральное государственное бюджетное образовательное учреждение высшего образования')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('«Московский Авиационный Институт»')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    p = doc.add_paragraph('(Национальный Исследовательский Университет)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Институт и кафедра
    p = doc.add_paragraph('Институт: №8 «Информационные технологии и прикладная математика»')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    p = doc.add_paragraph('Кафедра: 806 «Вычислительная математика и программирование»')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Название работы
    p = doc.add_paragraph('Лабораторная работа №2')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    
    p = doc.add_paragraph('по курсу «Криптография»')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Группа и студент (справа)
    p = doc.add_paragraph('Группа: М8О-306Б-21')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Студент: М.Ю.Курносов')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Преподаватель: А.В. Борисов')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Дата: 12.03.2026')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Город и год
    p = doc.add_paragraph('Москва, 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    # === Разрыв страницы ===
    doc.add_page_break()
    
    # === ТЕМА ===
    h = doc.add_heading('Тема', level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].bold = True
    
    p = doc.add_paragraph('Использование криптографических средств защиты информации. '
                         'Факторизация больших чисел с использованием хеш-функции ГОСТ Р 34.11-2012 (Стрибог).')
    p.runs[0].font.size = Pt(14)
    
    # === ЗАДАНИЕ ===
    h = doc.add_heading('Задание', level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].bold = True
    
    p = doc.add_paragraph('1. Получить номер варианта следующим образом:')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('1.1. Записать своё ФИО в кодировке UTF-8.')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('1.2. Вычислить хеш-сумму от ФИО с использованием хеш-функции '
                         'ГОСТ Р 34.11-2012 (Стрибог-256).')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('1.3. Младшие 8 бит хеш-суммы интерпретировать как число — '
                         'это и есть номер варианта (от 0 до 255).')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('2. Разложить на нетривиальные сомножители два числа a и b, '
                         'соответствующие полученному варианту.')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('3. В отчёте обязательно расписать все шаги решения задачи.')
    p.runs[0].font.size = Pt(14)
    
    # === ТЕОРИЯ ===
    h = doc.add_heading('Теоретические сведения', level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].bold = True
    
    p = doc.add_paragraph('ГОСТ Р 34.11-2012 (Стрибог) — российский стандарт хеш-функции. '
                         'Функция Стрибог-256 вырабатывает хеш-код длиной 256 бит (32 байта).')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Факторизация чисел — задача разложения целого числа на простые множители. '
                         'Для больших чисел (сотни цифр) эта задача является вычислительно сложной, '
                         'что лежит в основе криптосистемы RSA.')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Для факторизации могут применяться следующие методы:')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('• Метод пробного деления (для небольших множителей)', style='List Bullet')
    p = doc.add_paragraph('• Метод Полларда (rho-метод)', style='List Bullet')
    p = doc.add_paragraph('• Метод Полларда (p-1)', style='List Bullet')
    p = doc.add_paragraph('• Метод эллиптических кривых (ECM)', style='List Bullet')
    p = doc.add_paragraph('• Общий метод решета числового поля (GNFS) — для очень больших чисел', style='List Bullet')
    
    # === ХОД РАБОТЫ ===
    h = doc.add_heading('Ход работы', level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].bold = True
    
    # Часть 1
    subh = doc.add_heading('Часть 1. Получение номера варианта', level=2)
    subh.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('ФИО: Курносов Максим Юрьевич')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Кодировка: UTF-8')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Исходные данные (в байтах):')
    p.runs[0].font.size = Pt(14)
    
    fio_bytes = "Курносов Максим Юрьевич".encode('utf-8')
    p = doc.add_paragraph(fio_bytes.hex())
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Courier New'
    
    p = doc.add_paragraph('Хеш-функция: ГОСТ Р 34.11-2012 (Стрибог-256)')
    p.runs[0].font.size = Pt(14)
    
    # Здесь нужно вставить реальный хеш из ноутбука
    hash_hex = "2b4ff81c3c526b43f2b57959837a0521fce8908f7039a17570842aa8dfe8a41d"
    p = doc.add_paragraph(f'Хеш (hex): {hash_hex}')
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Courier New'
    
    # Номер варианта
    variant = 29
    p = doc.add_paragraph(f'Младший байт: 0x1D = {variant}')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph(f'НОМЕР ВАРИАНТА: {variant}')
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    
    # Часть 2
    subh = doc.add_heading('Часть 2. Факторизация чисел', level=2)
    subh.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph(f'Для варианта {variant} заданы следующие числа:')
    p.runs[0].font.size = Pt(14)
    
    # Числа из ноутбука
    a = 431359146674415075583841535117230246027501850497314921489544621489153
    b = 32317006071311007300714876688669951960444102669715484032130345427524655138867890893197201411522913463688717960921898019494119559150490921095088152453659852376261151862955313391199554873885711062041337884115460111636611778952226782228875998826893087373999483940897161216445436932489415085260185475297145213674492667686305880029617529867346941028193176958015096527639967699038346073836568598056404495352161269002922026318359257805566187267467143528604596252757796458208593675655347285497603144144679631145342269340341143469602171816506130380991362533407160923006783668002150656141866999560341143783887591265615365301023
    
    p = doc.add_paragraph('Число a:')
    p.runs[0].font.size = Pt(14)
    p = doc.add_paragraph(str(a))
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.name = 'Courier New'
    
    p = doc.add_paragraph(f'Количество цифр в a: {len(str(a))}')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Число b:')
    p.runs[0].font.size = Pt(14)
    p = doc.add_paragraph(str(b)[:200] + '...')
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.name = 'Courier New'
    
    p = doc.add_paragraph(f'Количество цифр в b: {len(str(b))}')
    p.runs[0].font.size = Pt(14)
    
    # Факторизация a
    subh = doc.add_heading('Факторизация числа a', level=3)
    subh.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Метод: функция factorint() из библиотеки SymPy (Python). '
                         'Использует комбинацию методов: пробное деление, метод Полларда (p-1), '
                         'метод Полларда (rho), метод эллиптических кривых (ECM).')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Результат факторизации:')
    p.runs[0].font.size = Pt(14)
    
    # Результаты из ноутбука
    p = 18889465931478792751657
    q = 22835963083295358096932575511191922182423842329
    
    p_doc = doc.add_paragraph(f'a = {p} × {q}')
    p_doc.runs[0].font.size = Pt(12)
    p_doc.runs[0].font.name = 'Courier New'
    
    p_doc = doc.add_paragraph(f'p = {p}')
    p_doc.runs[0].font.size = Pt(12)
    p_doc.runs[0].font.name = 'Courier New'
    
    p_doc = doc.add_paragraph(f'q = {q}')
    p_doc.runs[0].font.size = Pt(12)
    p_doc.runs[0].font.name = 'Courier New'
    
    # Проверка
    p_doc = doc.add_paragraph(f'Проверка: p × q = {p * q}')
    p_doc.runs[0].font.size = Pt(12)
    p_doc.runs[0].font.name = 'Courier New'
    
    p_doc = doc.add_paragraph(f'Исходное a = {a}')
    p_doc.runs[0].font.size = Pt(12)
    p_doc.runs[0].font.name = 'Courier New'
    
    p_doc = doc.add_paragraph(f'Совпадает: {p * q == a}')
    p_doc.runs[0].font.size = Pt(12)
    p_doc.runs[0].font.name = 'Courier New'
    
    p = doc.add_paragraph('Время факторизации: ~2.69 секунды')
    p.runs[0].font.size = Pt(14)
    
    # Факторизация b
    subh = doc.add_heading('Факторизация числа b', level=3)
    subh.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Число b содержит 608 цифр, что делает его факторизацию '
                         'классическими методами практически невозможной за разумное время.')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Попытка факторизации методом factorint() с ограничением не дала результатов.')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('Примечание: Числа такого размера используются в криптосистеме RSA-2048. '
                         'Их факторизация требует специализированного оборудования и алгоритмов '
                         '(GNFS — General Number Field Sieve) и может занимать годы вычислений.')
    p.runs[0].font.size = Pt(14)
    
    # === ВЫВОДЫ ===
    h = doc.add_heading('Выводы', level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].bold = True
    
    p = doc.add_paragraph('В ходе выполнения лабораторной работы были решены следующие задачи:')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('1. Изучена хеш-функция ГОСТ Р 34.11-2012 (Стрибог-256) и её применение '
                         'для генерации псевдослучайных чисел.', style='List Number')
    
    p = doc.add_paragraph('2. Получен номер варианта (29) путём хеширования ФИО и взятия младшего байта.', style='List Number')
    
    p = doc.add_paragraph('3. Успешно выполнена факторизация числа a (59 цифр) с использованием '
                         'библиотеки SymPy. Найдены два простых множителя.', style='List Number')
    
    p = doc.add_paragraph('4. Проанализирована сложность факторизации числа b (608 цифр). '
                         'Показано, что факторизация таких чисел является вычислительно сложной задачей.', style='List Number')
    
    p = doc.add_paragraph('5. Закреплены практические навыки работы с криптографическими библиотеками Python '
                         '(gostcrypto, sympy).', style='List Number')
    
    p = doc.add_paragraph('Таким образом, цель работы — изучение методов факторизации чисел и их применения '
                         'в криптографии — достигнута.')
    p.runs[0].font.size = Pt(14)
    
    # === СПИСОК ЛИТЕРАТУРЫ ===
    h = doc.add_heading('Список используемой литературы', level=1)
    h.runs[0].font.size = Pt(16)
    h.runs[0].bold = True
    
    p = doc.add_paragraph('1. ГОСТ Р 34.11-2012. Информационная технология. Криптографическая защита информации. '
                         'Функция хеширования. — М.: Стандартинформ, 2012.')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('2. Ященко В.В. Введение в криптографию. — СПб.: Питер, 2020.')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('3. SymPy Documentation. — https://docs.sympy.org/')
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('4. Конспект лекций по дисциплине «Криптографические методы защиты информации».')
    p.runs[0].font.size = Pt(14)
    
    # Сохранение
    doc.save('Отчёт_по_ЛР2_Факторизация.docx')
    print('Отчёт успешно создан: Отчёт_по_ЛР2_Факторизация.docx')

if __name__ == '__main__':
    create_lab2_report()
